import os
# 设置TESSDATA_PREFIX为tessdata的上一级目录
os.environ['TESSDATA_PREFIX'] = r'D:\soft\ocr'

import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from PIL import Image
import pytesseract
import pandas as pd

# 指定tesseract可执行文件路径
pytesseract.pytesseract.tesseract_cmd = r'D:\soft\ocr\tesseract.exe'

BASE_DIR = Path(__file__).resolve().parent.parent.parent
# 输入和输出文件路径
image_path = BASE_DIR / "datas/table.png"  # 替换为你的图片路径
pdf_path = BASE_DIR / "datas/有限空间管控流程图.pdf"
excel_path = BASE_DIR / "datas/output.xlsx"
word_path = BASE_DIR / "datas/output.docx"

output_folder = BASE_DIR / "datas/extracted_images"

def extract_images_from_pdf(pdf_path, output_folder):
    try:
        # 确保输出目录存在
        os.makedirs(output_folder, exist_ok=True)
        
        # 打开PDF文件
        pdf_document = fitz.open(pdf_path)
        
        # 遍历每一页
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # 获取页面中的图像列表
            image_list = page.get_images(full=True)
            
            # 遍历图像
            for img_index, img in enumerate(image_list):
                xref = img[0]  # 图像的xref编号
                base_image = pdf_document.extract_image(xref)  # 提取图像
                image_bytes = base_image["image"]  # 图像的字节数据
                image_ext = base_image["ext"]  # 图像扩展名（如jpeg, png）
                
                # 保存图像
                image_filename = f"{output_folder}/image_page{page_num+1}_{img_index}.{image_ext}"
                with open(image_filename, "wb") as image_file:
                    image_file.write(image_bytes)
                print(f"已保存图像：{image_filename}")
        
        # 关闭PDF
        pdf_document.close()
    except Exception as e:
        print(f"提取图像时发生错误: {e}")

def pdf_to_word(pdf_path, word_path):
    try:
        # 创建新的Word文档
        doc = Document()
        
        # 打开PDF文件
        pdf_document = fitz.open(pdf_path)
        
        # 遍历每一页
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # 提取文本
            text = page.get_text()
            
            # 将文本添加到Word文档
            doc.add_paragraph(text)
            
            # 提取图片
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # 保存临时图片
                temp_image_path = f"temp_image_{page_num}_{img_index}.{image_ext}"
                with open(temp_image_path, "wb") as image_file:
                    image_file.write(image_bytes)
                
                # 将图片添加到Word文档
                doc.add_picture(temp_image_path)
                
                # 删除临时图片
                os.remove(temp_image_path)
        
        # 保存Word文档
        doc.save(word_path)
        print(f"已成功将PDF转换为Word文档：{word_path}")
        
        # 关闭PDF
        pdf_document.close()
        return True
    except Exception as e:
        print(f"转换过程中发生错误: {e}")
        return False

# OCR图片转Word
def image_to_word(image_path, word_path):
    try:
        img = Image.open(image_path)
        custom_oem_psm_config = r'--tessdata-dir D:/soft/ocr/tessdata'
        text = pytesseract.image_to_string(img, lang='chi_sim', config=custom_oem_psm_config)
        doc = Document()
        doc.add_paragraph(text)
        doc.save(word_path)
        print(f"图片内容已写入Word：{word_path}")
    except Exception as e:
        print(f"图片转Word时发生错误: {e}\\n请检查chi_sim.traineddata是否在D:/soft/ocr/tessdata目录下，并确保TESSDATA_PREFIX和tesseract路径设置正确。")

# OCR图片转Excel（简单按行分列，适合结构化表格图片）
def image_to_excel(image_path, excel_path):
    try:
        img = Image.open(image_path)
        custom_oem_psm_config = r'--tessdata-dir D:/soft/ocr/tessdata'
        text = pytesseract.image_to_string(img, lang='chi_sim', config=custom_oem_psm_config)
        lines = text.split('\\n')
        data = [line.split() for line in lines if line.strip()]
        df = pd.DataFrame(data)
        df.to_excel(excel_path, index=False, header=False)
        print(f"图片内容已写入Excel：{excel_path}")
    except Exception as e:
        print(f"图片转Excel时发生错误: {e}\\n请检查chi_sim.traineddata是否在D:/soft/ocr/tessdata目录下，并确保TESSDATA_PREFIX和tesseract路径设置正确。")

# 主程序
def main():
    # 先提取图片
    extract_images_from_pdf(pdf_path, output_folder)
    # 假设只处理第一页第一个图片
    image_file = output_folder / "image_page1_0.png"
    image_to_word(image_file, word_path)
    image_to_excel(image_file, excel_path)

if __name__ == "__main__":
    main()